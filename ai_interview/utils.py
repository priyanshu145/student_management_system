import fitz          # PyMuPDF
import re
import os


# ─────────────────────────────────────────────────────────────
#  Public API
# ─────────────────────────────────────────────────────────────

def extract_resume_text(file_path: str) -> str:
    """
    Main entry point.  Detects file type and delegates to the
    appropriate extractor.  Returns cleaned plain text.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        raw = _extract_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        raw = _extract_docx(file_path)
    elif ext in ('.txt', '.md'):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw = f.read()
    else:
        raise ValueError(f"Unsupported resume format: {ext}")

    return _clean_text(raw)


# Keep the original function name for backwards compatibility
def extract_pdf_text(file_path: str) -> str:
    return extract_resume_text(file_path)


# ─────────────────────────────────────────────────────────────
#  PDF Extractor  (PyMuPDF — fast, handles multi-column)
# ─────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    text_parts = []
    try:
        doc = fitz.open(file_path)
        for page in doc:
            # "text" mode preserves reading order better than default
            text_parts.append(page.get_text("text"))
        doc.close()
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")
    return "\n".join(text_parts)


# ─────────────────────────────────────────────────────────────
#  DOCX Extractor  (python-docx — optional dependency)
# ─────────────────────────────────────────────────────────────

def _extract_docx(file_path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files. "
            "Install it with: pip install python-docx"
        )
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


# ─────────────────────────────────────────────────────────────
#  Text Cleaner
# ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalise whitespace, remove garbage characters, and trim.
    Keeps structure readable for the LLM.
    """
    if not text:
        return ""

    # Normalise unicode dashes / bullets to ASCII equivalents
    text = text.replace('\u2013', '-').replace('\u2014', '-')
    text = text.replace('\u2022', '*').replace('\u25cf', '*')
    text = text.replace('\u00a0', ' ')   # non-breaking space

    # Collapse repeated whitespace within a line
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line = re.sub(r'[ \t]{2,}', ' ', line).strip()
        # Skip lines that are just noise (lone punctuation, page numbers …)
        if re.fullmatch(r'[\s\-_=|•·.]{0,4}', line):
            continue
        cleaned.append(line)

    # Collapse 3+ consecutive blank lines into 2
    result = re.sub(r'\n{3,}', '\n\n', '\n'.join(cleaned))
    return result.strip()


# ─────────────────────────────────────────────────────────────
#  Optional: quick section parser (for richer prompting later)
# ─────────────────────────────────────────────────────────────

SECTION_HEADERS = [
    'experience', 'work experience', 'employment', 'education',
    'skills', 'technical skills', 'projects', 'certifications',
    'achievements', 'summary', 'objective', 'profile', 'contact',
]

def parse_resume_sections(text: str) -> dict:
    """
    Optionally split the resume into named sections.
    Returns a dict like {'skills': '...', 'experience': '...', ...}
    Useful if you want to pass targeted sections to the LLM.
    """
    sections = {}
    current_section = 'header'
    current_lines   = []

    for line in text.split('\n'):
        stripped = line.strip().lower().rstrip(':')
        if stripped in SECTION_HEADERS:
            # Save previous section
            if current_lines:
                sections[current_section] = '\n'.join(current_lines).strip()
            current_section = stripped
            current_lines   = []
        else:
            current_lines.append(line)

    # Don't forget last section
    if current_lines:
        sections[current_section] = '\n'.join(current_lines).strip()

    return sections